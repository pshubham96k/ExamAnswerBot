import os
import requests
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Function to read PDF and extract text
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + '\n'
    return text

# Response json to text 
def extract_text_from_response(api_response):
    """
    Extracts the 'text' part from the given API response.

    :param api_response: A dictionary containing the API response.
    :return: The extracted text if found, otherwise None.
    """
    try:
        candidates = api_response.get('candidates', [])
        if candidates:
            content = candidates[0].get('content', {})
            parts = content.get('parts', [])
            if parts:
                return parts[0].get('text', "").strip()
        return None
    except (KeyError, IndexError):
        return None

# Function to call Google Gemini API
def get_answers_from_gemini(prompt):
    api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key=AIzaSyAWCcJyC8xRozqpPDGEzL7KTRTFThkiRSY"
    headers = {
        "Content-Type": "application/json"
    }
    data = {
        "contents": [
            {
                "parts": [
                    {"text": prompt}
                ]
            }
        ]
    }
    try:
        response = requests.post(api_url, headers=headers, json=data)
        response.raise_for_status()
        response_data = response.json()
        print("API Response:", response_data)  # Debugging: print the full response
        text_data = extract_text_from_response(response_data)
        print(text_data)
        return text_data

    except requests.exceptions.RequestException as e:
        print(f"Error calling the Gemini API: {e}")
        return "Error generating answer."

# Function to create a Word document with questions and answers
def create_word_document(output_path, question_text, answer_text):
    doc = Document()
    
    # Add title
    title = doc.add_heading('Questions and Answers', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add question section
    question_heading = doc.add_heading('Question:', level=2)
    question_para = doc.add_paragraph()
    question_run = question_para.add_run(question_text)
    question_run.bold = False
    question_run.font.size = Pt(12)
    
    # Add some space between question and answer
    doc.add_paragraph()
    
    # Add answer section
    answer_heading = doc.add_heading('Answer:', level=2)
    answer_para = doc.add_paragraph()
    answer_run = answer_para.add_run(answer_text)
    answer_run.bold = False
    answer_run.font.size = Pt(12)
    
    # Add footer with line
    doc.add_paragraph('_' * 50)
    
    doc.save(output_path)

# Main function to process all exam papers
def process_exam_papers(input_folder):
    # Create output folder if it doesn't exist
    output_folder = os.path.join(input_folder, 'Processed_QA')
    os.makedirs(output_folder, exist_ok=True)
    
    for filename in os.listdir(input_folder):
        if filename.endswith('.pdf'):
            pdf_path = os.path.join(input_folder, filename)
            print(f"Processing {filename}...")
            
            # Extract question from PDF
            question_text = extract_text_from_pdf(pdf_path)
            print("Question:", question_text)
            
            # Get answer from Gemini API
            answer_text = get_answers_from_gemini(question_text)
            print("Answer:", answer_text)
            
            # Create output file in the new folder
            output_filename = f"{os.path.splitext(filename)[0]}_QA.docx"
            output_path = os.path.join(output_folder, output_filename)
            
            # Create Word document with both question and answer
            create_word_document(output_path, question_text, answer_text)
            print(f"Processed {filename} -> {output_filename}")

# Specify the folder containing the exam papers
input_folder = 'QuestionPaper'  # Change this to your folder path
process_exam_papers(input_folder)