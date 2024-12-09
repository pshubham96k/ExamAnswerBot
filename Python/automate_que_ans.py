import os
import re
import requests
import PyPDF2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_questions_from_pdf(pdf_path):
    """Extract questions from PDF file"""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + '\n'
            
        # Find questions using regex pattern
        # This pattern looks for numbered questions or questions ending with question marks
        questions = re.findall(r'(?:\d+[\)\.]\s*|\b[Qq]uestion\s*\d*[\:\.\)]\s*)(.*?\?)', text)
        
        # If no questions found with question marks, try to find numbered statements
        if not questions:
            questions = re.findall(r'\d+[\)\.]\s*(.*?(?:\?|\.|\n))', text)
            
        # Clean up the questions
        questions = [q.strip() for q in questions if q.strip()]
        
        return questions

def extract_text_from_response(api_response):
    """Extract text from API response"""
    try:
        candidates = api_response.get('candidates', [])
        if candidates:
            content = candidates[0].get('content', {})
            parts = content.get('parts', [])
            if parts:
                return parts[0].get('text', "").strip()
        return None
    except (KeyError, IndexError):
        return None

def get_answers_from_gemini(question):
    """Get answer from Gemini API for a single question"""
    api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key=AIzaSyAWCcJyC8xRozqpPDGEzL7KTRTFThkiRSY"
    headers = {
        "Content-Type": "application/json"
    }
    data = {
        "contents": [
            {
                "parts": [
                    {"text": question}
                ]
            }
        ]
    }
    try:
        response = requests.post(api_url, headers=headers, json=data)
        response.raise_for_status()
        response_data = response.json()
        return extract_text_from_response(response_data)
    except requests.exceptions.RequestException as e:
        print(f"Error calling the Gemini API: {e}")
        return "Error generating answer."

def create_qa_document(output_path, qa_pairs):
    """Create Word document with question-answer pairs"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Questions and Answers', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add each question-answer pair
    for i, (question, answer) in enumerate(qa_pairs, 1):
        # Add question
        q_heading = doc.add_heading(f'Question {i}:', level=2)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(question)
        q_run.font.size = Pt(12)
        
        # Add answer
        a_heading = doc.add_heading('Answer:', level=2)
        a_para = doc.add_paragraph()
        a_run = a_para.add_run(answer)
        a_run.font.size = Pt(12)
        
        # Add separator except for the last pair
        if i < len(qa_pairs):
            doc.add_paragraph('_' * 50)
    
    doc.save(output_path)

def process_exam_papers(input_folder):
    """Process all PDF files in the input folder"""
    # Create output folder if it doesn't exist
    output_folder = os.path.join(input_folder, 'QA_Output')
    os.makedirs(output_folder, exist_ok=True)
    
    for filename in os.listdir(input_folder):
        if filename.endswith('.pdf'):
            pdf_path = os.path.join(input_folder, filename)
            print(f"\nProcessing {filename}...")
            
            # Extract questions from PDF
            questions = extract_questions_from_pdf(pdf_path)
            
            if not questions:
                print(f"No questions found in {filename}")
                continue
                
            # Get answers for each question
            qa_pairs = []
            for question in questions:
                print(f"\nProcessing question: {question}")
                answer = get_answers_from_gemini(question)
                if answer:
                    qa_pairs.append((question, answer))
            
            if qa_pairs:
                # Create output file
                output_filename = f"{os.path.splitext(filename)[0]}_QA.docx"
                output_path = os.path.join(output_folder, output_filename)
                
                # Create Word document with Q&A pairs
                create_qa_document(output_path, qa_pairs)
                print(f"Created Q&A document: {output_filename}")
            else:
                print(f"No valid question-answer pairs generated for {filename}")

# Specify the folder containing the exam papers
input_folder = 'QuestionPaper'  # Change this to your folder path
process_exam_papers(input_folder)