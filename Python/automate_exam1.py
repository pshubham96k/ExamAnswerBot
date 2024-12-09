import os
import re
import requests
import PyPDF2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_text_from_pdf(pdf_path):
    """Extract complete text from PDF file"""
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text() + '\n'
        return text

def split_into_sections(text):
    """
    Split the text into sections based on questions.
    Each section contains the context and its associated question.
    """
    # This regex looks for numbered questions or questions ending with question marks
    # It preserves the text before each question as context
    sections = []
    
    # First, try to split by numbered questions
    question_pattern = r'(?:^|\n)(?:(?:\d+[\)\.]\s*|\b[Qq]uestion\s*\d*[\:\.\)]\s*).*?\?)'
    splits = re.split(f'({question_pattern})', text, flags=re.MULTILINE | re.DOTALL)
    
    current_context = ''
    current_question = ''
    
    for i, section in enumerate(splits):
        if not section.strip():
            continue
            
        # Check if this section is a question
        if re.match(question_pattern, section, re.MULTILINE):
            if current_question:
                # Store the previous section
                sections.append({
                    'context': current_context.strip(),
                    'question': current_question.strip()
                })
            current_question = section.strip()
            current_context = ''
        else:
            if current_question:
                # If we have a question, this is part of its context
                sections.append({
                    'context': current_context.strip(),
                    'question': current_question.strip()
                })
                current_question = ''
            current_context = section.strip()
    
    # Add the last section if there's a pending question
    if current_question:
        sections.append({
            'context': current_context.strip(),
            'question': current_question.strip()
        })
    
    return sections

def extract_text_from_response(api_response):
    """Extract text from API response"""
    try:
        candidates = api_response.get('candidates', [])
        if candidates:
            content = candidates[0].get('content', {})
            parts = content.get('parts', [])
            if parts:
                return parts[0].get('text', "").strip()
        return None
    except (KeyError, IndexError):
        return None

def get_answer_from_gemini(context, question):
    """Get answer from Gemini API including context"""
    api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key=AIzaSyAWCcJyC8xRozqpPDGEzL7KTRTFThkiRSY"
    headers = {
        "Content-Type": "application/json"
    }
    
    # Combine context and question for better answers
    prompt = f"{context}\n\nBased on the above context, please answer this question:\n{question}"
    
    data = {
        "contents": [
            {
                "parts": [
                    {"text": prompt}
                ]
            }
        ]
    }
    
    try:
        response = requests.post(api_url, headers=headers, json=data)
        response.raise_for_status()
        response_data = response.json()
        return extract_text_from_response(response_data)
    except requests.exceptions.RequestException as e:
        print(f"Error calling the Gemini API: {e}")
        return "Error generating answer."

def create_qa_document(output_path, sections_with_answers):
    """Create Word document with context, questions, and answers"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Questions and Answers', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add each section
    for i, section in enumerate(sections_with_answers, 1):
        # Add section number
        doc.add_heading(f'Section {i}', level=1)
        
        # Add context if it exists
        if section['context']:
            context_heading = doc.add_heading('Context:', level=2)
            context_para = doc.add_paragraph()
            context_run = context_para.add_run(section['context'])
            context_run.font.size = Pt(12)
        
        # Add question
        q_heading = doc.add_heading('Question:', level=2)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(section['question'])
        q_run.font.size = Pt(12)
        
        # Add answer
        a_heading = doc.add_heading('Answer:', level=2)
        a_para = doc.add_paragraph()
        a_run = a_para.add_run(section['answer'])
        a_run.font.size = Pt(12)
        
        # Add separator except for the last section
        if i < len(sections_with_answers):
            doc.add_paragraph('_' * 50)
    
    doc.save(output_path)

def process_exam_papers(input_folder):
    """Process all PDF files in the input folder"""
    # Create output folder if it doesn't exist
    output_folder = os.path.join(input_folder, 'QA_Output')
    os.makedirs(output_folder, exist_ok=True)
    
    for filename in os.listdir(input_folder):
        if filename.endswith('.pdf'):
            pdf_path = os.path.join(input_folder, filename)
            print(f"\nProcessing {filename}...")
            
            # Extract complete text from PDF
            text = extract_text_from_pdf(pdf_path)
            
            # Split into sections (context + questions)
            sections = split_into_sections(text)
            
            if not sections:
                print(f"No sections found in {filename}")
                continue
            
            # Get answers for each section
            sections_with_answers = []
            for section in sections:
                print(f"\nProcessing question: {section['question'][:100]}...")  # Show first 100 chars
                answer = get_answer_from_gemini(section['context'], section['question'])
                if answer:
                    section['answer'] = answer
                    sections_with_answers.append(section)
            
            if sections_with_answers:
                # Create output file
                output_filename = f"{os.path.splitext(filename)[0]}_QA.docx"
                output_path = os.path.join(output_folder, output_filename)
                
                # Create Word document with sections
                create_qa_document(output_path, sections_with_answers)
                print(f"Created Q&A document: {output_filename}")
            else:
                print(f"No valid sections with answers generated for {filename}")

# Specify the folder containing the exam papers
input_folder = 'QuestionPaper'  # Change this to your folder path
process_exam_papers(input_folder)